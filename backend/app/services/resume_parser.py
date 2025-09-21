import re
from pathlib import Path
from typing import Dict, List, Any, Optional
import logging

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Try importing the packages with fallbacks
try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    logger.warning("pdfplumber not found. Please install: pip install pdfplumber")
    PDFPLUMBER_AVAILABLE = False

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    logger.warning("python-docx not found. Please install: pip install python-docx")
    DOCX_AVAILABLE = False

try:
    import spacy
    nlp = spacy.load("en_core_web_sm")
    SPACY_AVAILABLE = True
except (ImportError, OSError):
    logger.warning("spaCy model not found. Using basic text processing.")
    SPACY_AVAILABLE = False
    nlp = None

async def parse_resume(file_path: Path) -> Dict[str, Any]:
    """Parse resume and extract structured data"""
    
    try:
        # Extract text based on file type
        if file_path.suffix.lower() == '.pdf':
            text = extract_text_from_pdf(file_path)
        elif file_path.suffix.lower() in ['.docx', '.doc']:
            text = extract_text_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_path.suffix}")
        
        if not text or len(text.strip()) < 50:
            raise ValueError("Resume text is too short or empty")
        
        # Extract structured data
        extracted_data = {
            "text": text,
            "word_count": len(text.split()),
            "skills": extract_skills(text),
            "experience": extract_experience_years(text),
            "education": extract_education(text),
            "projects": extract_projects(text),
            "certifications": extract_certifications(text),
            "contact_info": extract_contact_info(text),
            "languages": extract_languages(text),
            "achievements": extract_achievements(text)
        }
        
        logger.info(f"Successfully parsed resume: {file_path.name}")
        return extracted_data
        
    except Exception as e:
        logger.error(f"Error parsing resume {file_path}: {str(e)}")
        raise

def extract_text_from_pdf(file_path: Path) -> str:
    """Extract text from PDF using pdfplumber"""
    if not PDFPLUMBER_AVAILABLE:
        raise ImportError("pdfplumber is required for PDF processing. Install with: pip install pdfplumber")
    
    text = ""
    try:
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
                else:
                    logger.warning(f"No text found on page {page_num + 1} of {file_path.name}")
        
        if not text.strip():
            raise ValueError("No text could be extracted from PDF")
            
        return clean_text(text)
        
    except Exception as e:
        logger.error(f"Error extracting text from PDF {file_path}: {str(e)}")
        raise

def extract_text_from_docx(file_path: Path) -> str:
    """Extract text from DOCX file"""
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx is required for DOCX processing. Install with: pip install python-docx")
    
    text = ""
    try:
        doc = docx.Document(file_path)
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + "\n"
        
        if not text.strip():
            raise ValueError("No text could be extracted from DOCX")
            
        return clean_text(text)
        
    except Exception as e:
        logger.error(f"Error extracting text from DOCX {file_path}: {str(e)}")
        raise

def clean_text(text: str) -> str:
    """Clean and normalize text"""
    # Remove excessive whitespace
    text = re.sub(r'\s+', ' ', text)
    # Remove special characters but keep basic punctuation
    text = re.sub(r'[^\w\s\-\.\@\+\(\)\/\,\;]', ' ', text)
    # Remove multiple spaces
    text = re.sub(r' +', ' ', text)
    return text.strip()

def extract_skills(text: str) -> List[str]:
    """Extract technical and soft skills from resume text"""
    
    # Comprehensive skill database
    technical_skills = {
        # Programming Languages
        'python', 'java', 'javascript', 'typescript', 'c++', 'c#', 'php', 'ruby', 'go', 'rust',
        'swift', 'kotlin', 'scala', 'r', 'matlab', 'sql', 'nosql', 'html', 'css', 'scss', 'sass',
        
        # Frameworks & Libraries
        'react', 'angular', 'vue.js', 'vue', 'node.js', 'nodejs', 'express.js', 'express', 'django',
        'flask', 'fastapi', 'spring', 'spring boot', 'laravel', 'symfony', 'rails', 'ember.js',
        'backbone.js', 'jquery', 'bootstrap', 'tailwind', 'material-ui', 'antd',
        
        # Databases
        'mysql', 'postgresql', 'mongodb', 'redis', 'elasticsearch', 'cassandra', 'dynamodb',
        'sqlite', 'oracle', 'sql server', 'mariadb', 'couchdb',
        
        # Cloud & DevOps
        'aws', 'azure', 'gcp', 'google cloud', 'docker', 'kubernetes', 'jenkins', 'ci/cd',
        'terraform', 'ansible', 'puppet', 'chef', 'vagrant', 'heroku', 'vercel', 'netlify',
        
        # Data Science & AI/ML
        'machine learning', 'deep learning', 'ai', 'tensorflow', 'pytorch', 'keras', 'scikit-learn',
        'pandas', 'numpy', 'matplotlib', 'seaborn', 'plotly', 'jupyter', 'r studio', 'tableau',
        'power bi', 'spark', 'hadoop', 'kafka',
        
        # Mobile Development
        'react native', 'flutter', 'ios', 'android', 'swift', 'kotlin', 'xamarin', 'cordova',
        
        # Tools & Other
        'git', 'github', 'gitlab', 'bitbucket', 'jira', 'confluence', 'slack', 'trello',
        'postman', 'swagger', 'figma', 'sketch', 'adobe creative suite', 'photoshop'
    }
    
    soft_skills = {
        'leadership', 'teamwork', 'communication', 'problem solving', 'analytical thinking',
        'project management', 'time management', 'adaptability', 'creativity', 'collaboration',
        'critical thinking', 'decision making', 'negotiation', 'presentation', 'mentoring'
    }
    
    found_skills = []
    text_lower = text.lower()
    
    # Extract technical skills
    for skill in technical_skills:
        # Use word boundaries for better matching
        pattern = r'\b' + re.escape(skill) + r'\b'
        if re.search(pattern, text_lower):
            # Capitalize properly
            found_skills.append(skill.title())
    
    # Extract soft skills
    for skill in soft_skills:
        pattern = r'\b' + re.escape(skill) + r'\b'
        if re.search(pattern, text_lower):
            found_skills.append(skill.title())
    
    # Remove duplicates and limit results
    found_skills = list(set(found_skills))
    
    # Sort by relevance (technical skills first, then alphabetically)
    technical_found = [s for s in found_skills if s.lower() in technical_skills]
    soft_found = [s for s in found_skills if s.lower() in soft_skills]
    
    return sorted(technical_found) + sorted(soft_found)

def extract_experience_years(text: str) -> float:
    """Extract years of experience from resume text"""
    
    patterns = [
        # "5 years of experience", "3+ years experience", etc.
        r'(\d+(?:\.\d+)?)\s*[\+\-]?\s*(?:years?|yrs?)\s+(?:of\s+)?(?:experience|exp)',
        # "5+ years in", "3 years working", etc.
        r'(\d+(?:\.\d+)?)\s*[\+\-]?\s*(?:years?|yrs?)\s+(?:in|working|as|with)',
        # "experienced for 5 years"
        r'experienced\s+for\s+(\d+(?:\.\d+)?)\s*(?:years?|yrs?)',
        # "5 years professional experience"
        r'(\d+(?:\.\d+)?)\s*(?:years?|yrs?)\s+professional'
    ]
    
    years = []
    text_lower = text.lower()
    
    for pattern in patterns:
        matches = re.findall(pattern, text_lower)
        for match in matches:
            try:
                years.append(float(match))
            except ValueError:
                continue
    
    # If no explicit experience mentioned, try to infer from work history
    if not years:
        years = infer_experience_from_dates(text)
    
    return max(years) if years else 0.0

def infer_experience_from_dates(text: str) -> List[float]:
    """Infer experience from employment dates in resume"""
    
    # Look for date ranges like "2020-2023", "Jan 2020 - Present", etc.
    date_patterns = [
        r'(\d{4})\s*[-–]\s*(\d{4})',  # 2020-2023
        r'(\d{4})\s*[-–]\s*(?:present|current)',  # 2020-Present
    ]
    
    years = []
    current_year = 2024  # Update this to current year
    
    for pattern in date_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        for match in matches:
            try:
                start_year = int(match[0])
                if len(match) > 1 and match[1].isdigit():
                    end_year = int(match[1])
                else:
                    end_year = current_year  # Assume present
                
                experience = end_year - start_year
                if 0 < experience <= 50:  # Reasonable range
                    years.append(float(experience))
            except (ValueError, IndexError):
                continue
    
    return years

def extract_education(text: str) -> Dict[str, Any]:
    """Extract education information"""
    
    degree_patterns = [
        # Bachelor's degrees
        r'(?:bachelor|b\.?\s*(?:tech|sc|com|a|e|s)|b\.?tech|b\.?sc|b\.?com|b\.?a|b\.?e|b\.?s)',
        # Master's degrees  
        r'(?:master|m\.?\s*(?:tech|sc|com|a|e|s)|m\.?tech|m\.?sc|m\.?com|m\.?a|m\.?e|m\.?s|mba|ms)',
        # Doctoral degrees
        r'(?:doctor|ph\.?d|phd|doctorate)',
        # Other qualifications
        r'(?:diploma|certificate|associate)'
    ]
    
    degrees = []
    institutions = []
    graduation_years = []
    
    text_lower = text.lower()
    
    # Extract degrees
    for pattern in degree_patterns:
        matches = re.findall(pattern, text_lower)
        degrees.extend(matches)
    
    # Extract institutions (look for common university/college indicators)
    institution_patterns = [
        r'(?:university|college|institute|school)\s+of\s+[\w\s]+',
        r'[\w\s]+\s+(?:university|college|institute)',
        r'(?:iit|nit|bits|iiit|isi)\s*[\w\s]*'
    ]
    
    for pattern in institution_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        institutions.extend([m.strip() for m in matches])
    
    # Extract graduation years
    year_matches = re.findall(r'\b(19|20)\d{2}\b', text)
    for year in year_matches:
        year_int = int(year)
        if 1990 <= year_int <= 2030:  # Reasonable graduation year range
            graduation_years.append(year_int)
    
    return {
        "degrees": list(set([d.strip().title() for d in degrees]))[:5],
        "institutions": list(set(institutions))[:3],
        "graduation_years": sorted(set(graduation_years), reverse=True)[:3],
        "highest_degree": get_highest_degree(degrees)
    }

def get_highest_degree(degrees: List[str]) -> str:
    """Determine the highest degree from a list of degrees"""
    
    degree_hierarchy = {
        'phd': 5, 'doctorate': 5, 'doctor': 5,
        'master': 4, 'mba': 4, 'ms': 4, 'm.tech': 4, 'm.sc': 4,
        'bachelor': 3, 'b.tech': 3, 'b.sc': 3, 'b.com': 3,
        'diploma': 2,
        'certificate': 1
    }
    
    highest_level = 0
    highest_degree = "Unknown"
    
    for degree in degrees:
        degree_lower = degree.lower().strip()
        for key, level in degree_hierarchy.items():
            if key in degree_lower and level > highest_level:
                highest_level = level
                highest_degree = degree.strip().title()
    
    return highest_degree

def extract_projects(text: str) -> List[str]:
    """Extract project information from resume"""
    
    projects = []
    
    # Look for project sections
    project_keywords = ['project', 'projects:', 'key projects', 'major projects', 'notable projects']
    
    lines = text.split('\n')
    in_project_section = False
    current_project = ""
    
    for line in lines:
        line = line.strip()
        
        # Check if we're entering a project section
        if any(keyword in line.lower() for keyword in project_keywords):
            in_project_section = True
            continue
        
        # Check if we're leaving project section (new section header)
        if in_project_section and line and line[0].isupper() and ':' in line:
            if not any(keyword in line.lower() for keyword in ['project', 'work', 'experience']):
                in_project_section = False
                continue
        
        # Extract project information
        if in_project_section and line:
            # Look for bullet points or numbered items
            if re.match(r'^[\-\•\*\d\.]', line):
                if current_project:
                    projects.append(current_project.strip())
                current_project = re.sub(r'^[\-\•\*\d\.]+\s*', '', line)
            else:
                current_project += " " + line
        
        # Also look for project indicators throughout the resume
        elif any(indicator in line.lower() for indicator in ['developed', 'built', 'created', 'designed', 'implemented']):
            if 20 < len(line) < 200:  # Reasonable project description length
                projects.append(line.strip())
    
    # Add the last project if exists
    if current_project:
        projects.append(current_project.strip())
    
    # Clean and deduplicate projects
    cleaned_projects = []
    for project in projects:
        project = re.sub(r'\s+', ' ', project).strip()
        if len(project) > 20 and project not in cleaned_projects:
            cleaned_projects.append(project)
    
    return cleaned_projects[:6]  # Limit to 6 most relevant projects

def extract_certifications(text: str) -> List[str]:
    """Extract certifications and licenses"""
    
    certifications = []
    
    # Common certification patterns
    cert_patterns = [
        r'(?:certified|certification|certificate)\s+(?:in\s+)?([^\n]{5,80})',
        r'([^\n]*(?:aws|azure|google cloud|gcp)[^\n]*(?:certified|certification)[^\n]*)',
        r'([^\n]*(?:pmp|scrum|agile|cissp|ceh|cissp)[^\n]*)',
        r'([^\n]*(?:oracle|microsoft|cisco|comptia)[^\n]*(?:certified|certification)[^\n]*)'
    ]
    
    for pattern in cert_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        for match in matches:
            if isinstance(match, tuple):
                match = match[0] if match[0] else (match[1] if len(match) > 1 else "")
            
            cert = match.strip()
            if 5 < len(cert) < 100 and cert not in certifications:
                certifications.append(cert)
    
    return certifications[:8]  # Limit to 8 certifications

def extract_contact_info(text: str) -> Dict[str, str]:
    """Extract contact information"""
    
    # Email pattern
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    
    # Phone patterns
    phone_patterns = [
        r'[\+]?[1-9]?\d{1,4}?[-.\s]?\(?\d{1,3}?\)?[-.\s]?\d{1,4}[-.\s]?\d{1,4}[-.\s]?\d{1,9}',
        r'\+\d{1,3}[-.\s]?\d{1,4}[-.\s]?\d{1,4}[-.\s]?\d{4,10}',
        r'\(\d{3}\)\s?\d{3}[-.\s]?\d{4}'
    ]
    
    # LinkedIn pattern
    linkedin_pattern = r'(?:linkedin\.com/in/|linkedin\.com/pub/)([A-Za-z0-9\-_]+)'
    
    # GitHub pattern  
    github_pattern = r'(?:github\.com/)([A-Za-z0-9\-_]+)'
    
    emails = re.findall(email_pattern, text)
    
    phones = []
    for pattern in phone_patterns:
        matches = re.findall(pattern, text)
        phones.extend(matches)
    
    linkedin_matches = re.findall(linkedin_pattern, text, re.IGNORECASE)
    github_matches = re.findall(github_pattern, text, re.IGNORECASE)
    
    return {
        "email": emails[0] if emails else "",
        "phone": phones[0] if phones else "",
        "linkedin": linkedin_matches[0] if linkedin_matches else "",
        "github": github_matches[0] if github_matches else "",
        "portfolio": extract_portfolio_url(text)
    }

def extract_portfolio_url(text: str) -> str:
    """Extract portfolio/website URL"""
    
    # Look for common portfolio patterns
    url_patterns = [
        r'https?://(?:www\.)?[A-Za-z0-9\-_]+\.(?:com|net|org|io|dev|me)/[A-Za-z0-9\-_/]*',
        r'(?:portfolio|website):\s*(https?://[^\s]+)',
        r'(?:personal website|blog):\s*(https?://[^\s]+)'
    ]
    
    for pattern in url_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        if matches:
            return matches[0] if isinstance(matches[0], str) else matches[0][0]
    
    return ""

def extract_languages(text: str) -> List[str]:
    """Extract programming and spoken languages"""
    
    programming_languages = [
        'python', 'java', 'javascript', 'typescript', 'c++', 'c#', 'php', 'ruby',
        'go', 'rust', 'swift', 'kotlin', 'scala', 'r', 'matlab', 'sql'
    ]
    
    spoken_languages = [
        'english', 'spanish', 'french', 'german', 'chinese', 'japanese', 'korean',
        'hindi', 'arabic', 'portuguese', 'russian', 'italian', 'dutch'
    ]
    
    found_languages = []
    text_lower = text.lower()
    
    # Extract programming languages
    for lang in programming_languages:
        if re.search(r'\b' + re.escape(lang) + r'\b', text_lower):
            found_languages.append(f"{lang.title()} (Programming)")
    
    # Extract spoken languages
    for lang in spoken_languages:
        if re.search(r'\b' + re.escape(lang) + r'\b', text_lower):
            found_languages.append(f"{lang.title()} (Spoken)")
    
    return found_languages

def extract_achievements(text: str) -> List[str]:
    """Extract achievements and awards"""
    
    achievements = []
    achievement_keywords = [
        'achievement', 'award', 'recognition', 'honor', 'accomplishment',
        'winner', 'first place', 'second place', 'third place', 'medal',
        'scholarship', 'dean\'s list', 'magna cum laude', 'summa cum laude'
    ]
    
    lines = text.split('\n')
    
    for line in lines:
        line = line.strip()
        if any(keyword in line.lower() for keyword in achievement_keywords):
            if 10 < len(line) < 150:  # Reasonable achievement description length
                achievements.append(line)
    
    return achievements[:5]  # Limit to 5 achievements